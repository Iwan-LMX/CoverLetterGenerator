"""Tools for web scraping and file handling."""

import requests
from bs4 import BeautifulSoup
from urllib.parse import urlparse
import re
from pathlib import Path
from typing import Optional, Dict, Any
from .config import Settings


class WebScraper:
    """Tool for scraping job descriptions from web pages."""
    
    def __init__(self):
        self.session = requests.Session()
        self.session.headers.update({'User-Agent': Settings.USER_AGENT})
    
    def scrape_job_description(self, url: str) -> Dict[str, Any]:
        """Scrape job description from a URL.
        
        Args:
            url: URL of the job posting
            
        Returns:
            Dictionary containing job information
        """
        try:
            response = self.session.get(url, timeout=Settings.REQUEST_TIMEOUT)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Extract basic info
            title = self._extract_title(soup)
            company = self._extract_company(soup, url)
            description = self._extract_description(soup)
            
            return {
                'title': title,
                'company': company,
                'description': description,
                'url': url,
                'raw_html': str(soup)
            }
            
        except requests.RequestException as e:
            raise Exception(f"Error fetching URL {url}: {str(e)}")
        except Exception as e:
            raise Exception(f"Error parsing job description: {str(e)}")
    
    def _extract_title(self, soup: BeautifulSoup) -> str:
        """Extract job title from the page."""
        # First check meta tags
        meta_title = soup.select_one('meta[property="og:title"]')
        if meta_title and meta_title.get('content'):
            title = meta_title.get('content').strip()
            if title and title != "Job":
                return title
        
        # Extended selectors for job titles
        selectors = [
            'h1',
            '[data-ui="job-title"]',
            '.job-title',
            '.jobsearch-JobInfoHeader-title',
            '[data-testid="job-title"]',
            '.job-header-title',
            '.position-title',
            '.job-name',
            '.jobs-unified-top-card__job-title',
            '.styles__JobTitle',
            '[data-automation-id="jobTitle"]'
        ]
        
        for selector in selectors:
            element = soup.select_one(selector)
            if element:
                title = element.get_text(strip=True)
                if title and len(title) > 2:  # Must be substantial
                    return title
        
        # Check JSON-LD structured data
        json_title = self._extract_title_from_json_ld(soup)
        if json_title:
            return json_title
        
        # Fallback to page title (clean it up)
        title_tag = soup.find('title')
        if title_tag:
            title = title_tag.get_text(strip=True)
            # Remove common suffixes
            for suffix in [" - Jobs", " | Jobs", " - Careers", " | Careers", " - Apply Now"]:
                title = title.replace(suffix, "")
            return title
        
        return "Unknown Position"
    
    def _extract_company(self, soup: BeautifulSoup, url: str) -> str:
        """Extract company name from the page."""
        # Extended selectors for company names
        selectors = [
            '[data-ui="company-name"]',
            '.company-name',
            '.jobsearch-InlineCompanyRating-companyHeader',
            '[data-testid="company-name"]',
            '.company',
            '.employer-name',
            '.jobs-unified-top-card__company-name',
            '.styles__CompanyName',
            '[data-automation-id="companyName"]',
            '.company-header__name',
            'a[data-cy="company-name"]'
        ]
        
        for selector in selectors:
            element = soup.select_one(selector)
            if element:
                company = element.get_text(strip=True)
                if company and len(company) > 1:
                    return company
        
        # Check JSON-LD structured data
        json_company = self._extract_company_from_json_ld(soup)
        if json_company:
            return json_company
        
        # Check meta tags
        meta_site = soup.select_one('meta[property="og:site_name"]')
        if meta_site and meta_site.get('content'):
            return meta_site.get('content').strip()
        
        # Fallback to domain name (improved)
        domain = urlparse(url).netloc
        company_name = domain.replace('www.', '').split('.')[0]
        
        # Handle common job board domains
        job_boards = {
            'apply': 'Workable Application',
            'workable': 'Workable',
            'indeed': 'Indeed',
            'linkedin': 'LinkedIn',
            'glassdoor': 'Glassdoor',
            'angellist': 'AngelList',
            'jobs': 'Job Board'
        }
        
        return job_boards.get(company_name.lower(), company_name.title())
    
    def _extract_description(self, soup: BeautifulSoup) -> str:
        """Extract job description from the page."""
        # First, try to get description from meta tags (works for many JS-heavy sites)
        meta_description = self._extract_meta_description(soup)
        if meta_description and len(meta_description) > 100:  # Must be substantial
            return self._clean_text(meta_description)
        
        # Remove script and style elements
        for script in soup(["script", "style", "nav", "header", "footer"]):
            script.decompose()
        
        # Extended selectors for job descriptions (including Workable and other common sites)
        selectors = [
            # Generic selectors
            '.job-description',
            '.jobsearch-jobDescriptionText',
            '[data-testid="job-description"]',
            '.description',
            '.job-content',
            '.job-details',
            '.position-description',
            
            # Workable specific
            '[data-ui="job-description"]',
            '.styles__JobDescription',
            '.job-ad-description',
            
            # LinkedIn
            '.show-more-less-html__markup',
            '.jobs-description__content',
            
            # Indeed
            '#jobDescriptionText',
            '.jobsearch-jobDescriptionText',
            
            # Glassdoor
            '.jobDescriptionContent',
            
            # AngelList/Wellfound
            '.job-description-text',
            
            # Generic content areas
            '[role="main"]',
            '.main-content',
            '.content',
            'main',
            'article'
        ]
        
        for selector in selectors:
            element = soup.select_one(selector)
            if element:
                text = element.get_text(separator=' ', strip=True)
                cleaned_text = self._clean_text(text)
                if len(cleaned_text) > 50:  # Must have substantial content
                    return cleaned_text
        
        # If still no content, try to extract from JSON-LD structured data
        json_ld_desc = self._extract_from_json_ld(soup)
        if json_ld_desc:
            return self._clean_text(json_ld_desc)
        
        # Final fallback: try to get meaningful content from body
        body = soup.find('body')
        if body:
            # Look for any paragraphs or divs with substantial text
            content_elements = body.find_all(['p', 'div'], string=True)
            combined_text = ""
            for elem in content_elements:
                text = elem.get_text(strip=True)
                if len(text) > 30:  # Only include substantial text blocks
                    combined_text += text + " "
            
            if combined_text:
                cleaned = self._clean_text(combined_text)
                return cleaned[:2000] if len(cleaned) > 2000 else cleaned
        
        # If we still have nothing, return the meta description or a message
        return meta_description if meta_description else "Could not extract job description. The page may require JavaScript to load content."
    
    def _extract_meta_description(self, soup: BeautifulSoup) -> str:
        """Extract description from meta tags."""
        meta_selectors = [
            'meta[name="description"]',
            'meta[property="og:description"]',
            'meta[name="twitter:description"]'
        ]
        
        for selector in meta_selectors:
            meta = soup.select_one(selector)
            if meta and meta.get('content'):
                return meta.get('content')
        
        return ""
    
    def _extract_from_json_ld(self, soup: BeautifulSoup) -> str:
        """Extract job description from JSON-LD structured data."""
        scripts = soup.find_all('script', type='application/ld+json')
        for script in scripts:
            try:
                import json
                data = json.loads(script.string)
                
                # Handle both single objects and arrays
                if isinstance(data, list):
                    data = data[0] if data else {}
                
                # Look for JobPosting schema
                if data.get('@type') == 'JobPosting':
                    desc = data.get('description', '')
                    if desc and len(desc) > 50:
                        return desc
                        
            except (json.JSONDecodeError, AttributeError):
                continue
        
        return ""
    
    def _extract_title_from_json_ld(self, soup: BeautifulSoup) -> str:
        """Extract job title from JSON-LD structured data."""
        scripts = soup.find_all('script', type='application/ld+json')
        for script in scripts:
            try:
                import json
                data = json.loads(script.string)
                
                if isinstance(data, list):
                    data = data[0] if data else {}
                
                if data.get('@type') == 'JobPosting':
                    title = data.get('title', '')
                    if title:
                        return title
                        
            except (json.JSONDecodeError, AttributeError):
                continue
        
        return ""
    
    def _extract_company_from_json_ld(self, soup: BeautifulSoup) -> str:
        """Extract company name from JSON-LD structured data."""
        scripts = soup.find_all('script', type='application/ld+json')
        for script in scripts:
            try:
                import json
                data = json.loads(script.string)
                
                if isinstance(data, list):
                    data = data[0] if data else {}
                
                if data.get('@type') == 'JobPosting':
                    hiring_org = data.get('hiringOrganization', {})
                    if isinstance(hiring_org, dict):
                        name = hiring_org.get('name', '')
                        if name:
                            return name
                        
            except (json.JSONDecodeError, AttributeError):
                continue
        
        return ""
    
    def _clean_text(self, text: str) -> str:
        """Clean extracted text."""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters that might cause issues
        text = re.sub(r'[^\w\s\.,!?;:()\-\'/"]', '', text)
        return text.strip()


class FileHandler:
    """Tool for handling file operations."""
    
    @staticmethod
    def read_template(template_path: Optional[Path] = None) -> str:
        """Read the cover letter template.
        
        Args:
            template_path: Path to template file. If None, uses default.
            
        Returns:
            Template content
        """
        path = template_path or Settings.TEMPLATE_PATH
        
        if not path.exists():
            # Create default template if it doesn't exist
            FileHandler.create_default_template(path)
        
        return path.read_text(encoding='utf-8')
    
    @staticmethod
    def save_cover_letter(content: str, filename: Optional[str] = None, 
                         company_name: Optional[str] = None, 
                         position_title: Optional[str] = None) -> Dict[str, Path]:
        """Save the generated cover letter to both PDF and TEXT files in a dedicated subfolder.
        
        Args:
            content: Cover letter content
            filename: Base filename without extension. If None, generates timestamp-based name.
            company_name: Company name for folder organization (optional)
            position_title: Position title for folder organization (optional)
            
        Returns:
            Dictionary with 'txt' and 'pdf' paths to the saved files
        """
        from datetime import datetime
        
        # Generate base filename if not provided
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            if company_name and position_title:
                # Create a clean filename from company and position
                clean_company = FileHandler._clean_filename(company_name)
                clean_position = FileHandler._clean_filename(position_title)
                filename = f"cover_letter_{clean_company}_{clean_position}_{timestamp}"
            else:
                filename = f"cover_letter_{timestamp}"
        else:
            # Remove extension if provided
            filename = Path(filename).stem
        
        # Create subfolder for this cover letter
        output_base_dir = Settings.ensure_output_dir()
        subfolder_name = f"{filename}"
        output_dir = output_base_dir / subfolder_name
        output_dir.mkdir(exist_ok=True)
        
        # Save TEXT file
        txt_path = output_dir / f"{filename}.txt"
        txt_path.write_text(content, encoding='utf-8')
        
        # Save PDF file
        pdf_path = output_dir / f"{filename}.pdf"
        FileHandler._save_as_pdf(content, pdf_path)
        
        return {
            'txt': txt_path,
            'pdf': pdf_path,
            'folder': output_dir
        }
    
    @staticmethod
    def create_default_template(template_path: Path):
        """Create a default cover letter template."""
        default_template = """Dear Hiring Manager,

I am writing to express my strong interest in the {position} position at {company}. With my background in {relevant_experience}, I am excited about the opportunity to contribute to your team.

{body_paragraph_1}

{body_paragraph_2}

Thank you for considering my application. I look forward to hearing from you soon.

Sincerely,
[Your Name]"""
        
        template_path.parent.mkdir(parents=True, exist_ok=True)
        template_path.write_text(default_template, encoding='utf-8')
    
    @staticmethod
    def _clean_filename(text: str) -> str:
        """Clean text for use as filename.
        
        Args:
            text: Text to clean
            
        Returns:
            Cleaned filename-safe string
        """
        # Remove or replace invalid filename characters
        cleaned = re.sub(r'[<>:"/\\|?*]', '', text)
        # Replace spaces and special chars with underscores
        cleaned = re.sub(r'[\s\-\.\,\(\)]+', '_', cleaned)
        # Remove multiple underscores
        cleaned = re.sub(r'_+', '_', cleaned)
        # Trim underscores and limit length
        cleaned = cleaned.strip('_')[:50]
        
        return cleaned if cleaned else 'unnamed'
    
    @staticmethod
    def _save_as_pdf(content: str, pdf_path: Path):
        """Save content as PDF file.
        
        Args:
            content: Text content to save as PDF
            pdf_path: Path where to save the PDF file
        """
        try:
            # Try using reportlab (preferred method)
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            
            # Create PDF document
            doc = SimpleDocTemplate(str(pdf_path), pagesize=letter,
                                  rightMargin=72, leftMargin=72,
                                  topMargin=72, bottomMargin=18)
            
            # Get styles
            styles = getSampleStyleSheet()
            
            # Create custom style for cover letter
            cover_letter_style = ParagraphStyle(
                'CoverLetter',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=12,
                leading=14
            )
            
            # Split content into paragraphs and create flowables
            story = []
            paragraphs = content.split('\n\n')
            
            for para in paragraphs:
                if para.strip():
                    # Clean paragraph text for reportlab
                    clean_para = para.strip().replace('\n', '<br/>')
                    story.append(Paragraph(clean_para, cover_letter_style))
                    story.append(Spacer(1, 6))
            
            # Build PDF
            doc.build(story)
            
        except ImportError:
            # Fallback: try using weasyprint
            try:
                import weasyprint
                
                # Create HTML content
                # Process content outside f-string to avoid backslash issues
                processed_content = '<p>' + content.replace('\n\n', '</p><p>').replace('\n', '<br>') + '</p>'
                
                html_content = f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <meta charset="UTF-8">
                    <style>
                        body {{
                            font-family: 'Times New Roman', Times, serif;
                            font-size: 11pt;
                            line-height: 1.4;
                            margin: 1in;
                            max-width: 8.5in;
                        }}
                        p {{
                            margin-bottom: 12pt;
                        }}
                    </style>
                </head>
                <body>
                    {processed_content}
                </body>
                </html>
                """
                
                weasyprint.HTML(string=html_content).write_pdf(str(pdf_path))
                
            except ImportError:
                # Final fallback: use fpdf (simpler but available)
                try:
                    from fpdf import FPDF
                    
                    pdf = FPDF()
                    pdf.add_page()
                    pdf.set_font('Times', size=11)
                    
                    # Add content line by line
                    lines = content.split('\n')
                    for line in lines:
                        if line.strip():
                            # Handle long lines by wrapping
                            words = line.split(' ')
                            current_line = ''
                            for word in words:
                                test_line = current_line + (' ' if current_line else '') + word
                                # Approximate character limit per line
                                if len(test_line) <= 80:
                                    current_line = test_line
                                else:
                                    if current_line:
                                        pdf.cell(0, 6, current_line.encode('latin-1', 'replace').decode('latin-1'), ln=True)
                                    current_line = word
                            
                            if current_line:
                                pdf.cell(0, 6, current_line.encode('latin-1', 'replace').decode('latin-1'), ln=True)
                        else:
                            pdf.ln(6)  # Empty line
                    
                    pdf.output(str(pdf_path))
                    
                except ImportError:
                    # If no PDF library is available, create a simple text file with PDF extension
                    # and add a note about PDF generation
                    fallback_content = f"""# Cover Letter (PDF generation failed - install reportlab, weasyprint, or fpdf2)

{content}

---
Note: To generate proper PDF files, install one of these packages:
- pip install reportlab (recommended)
- pip install weasyprint
- pip install fpdf2
"""
                    pdf_path.write_text(fallback_content, encoding='utf-8')
                    print(f"⚠️  PDF libraries not available. Created text file with PDF extension at {pdf_path}")
                    
        except Exception as e:
            # If PDF generation fails, create a fallback text file
            fallback_content = f"""# Cover Letter (PDF generation error: {str(e)})

{content}
"""
            pdf_path.write_text(fallback_content, encoding='utf-8')
            print(f"⚠️  PDF generation failed: {e}. Created text fallback at {pdf_path}")
    
    @staticmethod
    def parse_resume_file(file_path: Path) -> str:
        """Parse resume from various file formats.
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            Extracted text content from the resume
            
        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file doesn't exist
        """
        if not file_path.exists():
            raise FileNotFoundError(f"Resume file not found: {file_path}")
        
        file_extension = file_path.suffix.lower()
        
        try:
            if file_extension == '.pdf':
                return FileHandler._extract_text_from_pdf(file_path)
            elif file_extension == '.docx':
                return FileHandler._extract_text_from_docx(file_path)
            elif file_extension in ['.txt', '.md']:
                return file_path.read_text(encoding='utf-8')
            else:
                raise ValueError(f"Unsupported file format: {file_extension}. Supported formats: .pdf, .docx, .txt, .md")
                
        except Exception as e:
            raise Exception(f"Error parsing resume file {file_path}: {str(e)}")
    
    @staticmethod
    def _extract_text_from_pdf(pdf_path: Path) -> str:
        """Extract text from PDF file."""
        try:
            import PyPDF2
            
            text = ""
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            if not text.strip():
                raise Exception("No text could be extracted from the PDF. The PDF might be image-based or encrypted.")
            
            return text.strip()
            
        except ImportError:
            raise Exception("PyPDF2 is required to read PDF files. Install it with: pip install PyPDF2")
        except Exception as e:
            # Try alternative PDF parsing method
            try:
                return FileHandler._extract_text_from_pdf_alternative(pdf_path)
            except:
                raise Exception(f"Could not extract text from PDF: {str(e)}")
    
    @staticmethod
    def _extract_text_from_pdf_alternative(pdf_path: Path) -> str:
        """Alternative PDF text extraction using pdfplumber (if available)."""
        try:
            import pdfplumber
            
            text = ""
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            return text.strip()
            
        except ImportError:
            raise Exception("Could not extract text from PDF. Install pdfplumber for better PDF support: pip install pdfplumber")
    
    @staticmethod
    def _extract_text_from_docx(docx_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            from docx import Document
            
            doc = Document(docx_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            if not text.strip():
                raise Exception("No text found in the DOCX file.")
            
            return text.strip()
            
        except ImportError:
            raise Exception("python-docx is required to read DOCX files. Install it with: pip install python-docx")
        except Exception as e:
            raise Exception(f"Could not extract text from DOCX file: {str(e)}")